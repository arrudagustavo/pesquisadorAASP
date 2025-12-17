import os
import re
import json
import time
import zipfile
import traceback
import gc
from flask import Flask, render_template, request, send_file, flash, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt 
from docx.enum.text import WD_COLOR_INDEX

app = Flask(__name__)
app.secret_key = 'segredo_juridico_absoluto'

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

current_status = {"message": "Aguardando...", "percent": 0}

def update_status(msg, percent):
    global current_status
    current_status["message"] = msg
    current_status["percent"] = percent

# ==============================================================================
# 1. VALIDADOR V75
# ==============================================================================
def validate_final_term(text):
    if not text: return False
    t = text.strip()
    u = t.upper()
    
    if len(t) < 3: return False
    if len(t) > 3000: return False 
    
    # OAB DEATH TOUCH
    if re.search(r'\bO\.?A\.?B\b', u): return False
    if re.search(r'[A-Z]{2}\s*[\-]\s*\d{4,}', u): return False 

    # Pontuação inicial
    if t.startswith(('/', '-', '.', ',', ':')): return False
    # Siglas soltas
    if re.match(r'^[\W_]*[A-Z]{2}[\W_]*$', u): return False

    # Números soltos
    if re.match(r'^[0-9\-\.\/\s]+$', t): return False 
    
    # Datas
    if re.search(r'\d{1,2}\s+DE\s+[A-ZÇ]+\s+DE\s+\d{4}', u): return False

    BLOCK_PHRASES = [
        "GERADO EM", "ASSOCIADO:", "ASSOCIADO",
        "D J E N", "TJRJ", "TJMG", "TJSP", "STJ", "TRT", "TJPR", "TRF1",
        "DISPONIBILIZAÇÃO", "PUBLICAÇÃO", "ARQUIVO:",
        "DIÁRIO ELETRÔNICO", "DIÁRIO DA JUSTIÇA",
        "TIPO DE COMUNICAÇÃO", "MEIO:", "DATA DE",
        "PODER JUDICIÁRIO", "JUSTIÇA DE PRIMEIRA",
        "MINISTÉRIO PÚBLICO", "DEFENSORIA PÚBLICA",
        "ADMINISTRADORA JUDICIAL", "REPRESENTANTE DO MINISTÉRIO", "ADMINISTRAÇÃO JUDICIAL",
        "VARA CÍVEL", "VARA EMPRESARIAL", "COMARCA DE", "CARTÓRIO", "CÂMARA CÍVEL", "CÂMARA",
        "JUIZ DE DIREITO", "ESCRIVÃ", "DIRETOR DE SECRETARIA", 
        "RELATOR", "RELATOR:", "AGRAVO DE INSTRUMENTO", "AGRAVO", "LIMINAR",
        "RECURSO", "ASSUNTO", "ASSUNTO:", "DESPACHOS", "SECRETARIA DA",
        "CENTRAL PARA PROCEDER", "ID ", "AUTOS", "FLS.", "ADVOGADO:", "ADVOGADO",
        "QUADRO GERAL DE CREDORES", "HOMOLOGADO", 
        "HABILITAÇÃO DE CRÉDITO", "RETARDATÁRIA",
        "EM RECUPERAÇÃO JUDICIAL", "RECUPERAÇÃO JUDICIAL", "MASSA FALIDA",
        "DECISÃO", "SENTENÇA", "INTIME-SE", "PUBLIQUE-SE",
        "PREPOTENTE", "OPORTUNISTA", "AUTORITÁRIA", "GARRAS",
        "UMBILICALMENTE", "CAPACIDADE OPERACIONAL", "FLUXO DE CAIXA",
        "AUTOTUTELA", "PIRATAS RECUPERACIONAIS", "ODIOSA", "CONSIDERANDO"
    ]

    for block in BLOCK_PHRASES:
        if block in u: return False

    EXACT_BLOCKS = ["DECISÃO", "SENTENÇA", "DESPACHO", "VISTOS", "ÓRGÃO:", "ADVOGADO(S)", "INTIMAÇÃO", "PARTE(S):", "ADVOGADOS", "ADVOGADO", "RECORRENTE", "RECORRIDO", "AGRAVANTE", "AGRAVADO", "AUTOR", "RÉU", "EXECUTADO", "EXEQUENTE", "IMPETRANTE", "IMPETRADO", "IMPUGNANTE", "IMPUGNADO", "REQUERENTE", "REQUERIDO", "POLO ATIVO", "POLO PASSIVO", "POLO", "PASSIVO"]
    if u in EXACT_BLOCKS: return False
        
    return True

# ==============================================================================
# 2. GUARDA DE SINTAXE
# ==============================================================================
def is_valid_start(text):
    if not text: return False
    
    # Palavras proibidas no INÍCIO do nome (Conectivos)
    forbidden = [
        "para", "que", "de", "do", "da", "dos", "das", 
        "e", "a", "à", "ao", "aos", "o", "os", "as",
        "na", "no", "nas", "nos", "ne", "nesta", "neste",
        "em", "com", "por", "pelo", "pela", "pelos", "pelas",
        "sob", "sobre", "tal", "como", "quando", "se", "houver", "onde", 
        "outro", "outra", "outras", "to"
    ]
    
    first_word = text.strip().split(' ')[0].lower()
    first_word = re.sub(r'[^\w]', '', first_word)
    
    if first_word in forbidden:
        return False
    
    return True

# ==============================================================================
# 3. SANITIZAÇÃO
# ==============================================================================
def sanitize_docx_xml(filepath):
    try:
        with zipfile.ZipFile(filepath, 'r') as zin:
            xml_content = zin.read('word/document.xml').decode('utf-8')
        xml_content = re.sub(r'<w:proofErr[^>]*/>', '', xml_content)
        xml_content = re.sub(r'<w:noBreakHyphen[^>]*/>', '', xml_content)
        xml_content = re.sub(r'<w:softHyphen[^>]*/>', '', xml_content)
        xml_content = re.sub(r'<w:lastRenderedPageBreak[^>]*/>', '', xml_content)
        new_filepath = filepath.replace('.docx', '_clean.docx')
        with zipfile.ZipFile(filepath, 'r') as zin:
            with zipfile.ZipFile(new_filepath, 'w') as zout:
                for item in zin.infolist():
                    if item.filename == 'word/document.xml':
                        zout.writestr(item, xml_content)
                    else:
                        zout.writestr(item, zin.read(item.filename))
        return new_filepath
    except: return filepath

# ==============================================================================
# 4. EXTRATOR V75 (HYPHEN HANDLER)
# ==============================================================================
def process_document_by_publication(doc):
    all_paras = []
    all_paras.extend(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
                
    total_count = 0
    current_pub_text = []
    current_pub_paras = []
    
    new_pub_marker = re.compile(r'^\s*\d+\.\s+[A-Z]', re.IGNORECASE)

    for i, para in enumerate(all_paras):
        if i % 100 == 0: gc.collect()
        text = para.text.strip()
        
        if text and new_pub_marker.match(text):
            if current_pub_paras:
                total_count += process_single_publication(current_pub_text, current_pub_paras)
            current_pub_text = [text]
            current_pub_paras = [para]
        else:
            if text: current_pub_text.append(text)
            current_pub_paras.append(para)

    if current_pub_paras:
        total_count += process_single_publication(current_pub_text, current_pub_paras)

    return total_count

def process_single_publication(text_lines, paragraphs):
    # 1. Normalização Total
    full_text = " ".join(text_lines).replace('\n', ' ').replace('\r', ' ')
    while "  " in full_text: full_text = full_text.replace("  ", " ")
    
    # 2. Remoção de Lixo Gráfico
    full_text = re.sub(r'[^\w\s\.\-\/\(\),:]', ' ', full_text)

    # 3. Injeção de Espaços (Anti-Cola)
    keywords = ["Advogado", "Advogada", "OAB", "Relator", "Juiz", "Poder Judiciário", "Intimação", "Processo:", "Adv:", "Adv -"]
    for kw in keywords:
        full_text = re.sub(rf'(\S)({kw})', r'\1 \2', full_text, flags=re.IGNORECASE)

    # 4. NORMALIZAÇÃO DE CABEÇALHOS E PAPÉIS
    
    # MK_PARTES
    regex_polos = r'(?:Parte\s*\(?s\)?|Polo\s+Ativo|Polo\s+Passivo|Polo|Passivo|Destinatário\s*\(?s\)?)\s*:'
    full_text = re.sub(regex_polos, 'MK_PARTES:', full_text, flags=re.IGNORECASE)
    
    # === NORMALIZAÇÃO COM HIFEN OU DOIS PONTOS ===
    # O padrão [:\-] aceita tanto dois pontos quanto hífen como separador
    
    # Impugnante/Impugnado
    full_text = re.sub(r'Impugnante\s*[:\-]', 'Impugnante:', full_text, flags=re.IGNORECASE)
    full_text = re.sub(r'Impugnad[oa]\s*\(?a?\)?\s*[:\-]', 'Impugnado:', full_text, flags=re.IGNORECASE)
    
    # Requerente/Requerido
    full_text = re.sub(r'Requerente\s*[:\-]', 'Requerente:', full_text, flags=re.IGNORECASE)
    full_text = re.sub(r'Requerid[oa]\s*\(?a?\)?\s*[:\-]', 'Requerido:', full_text, flags=re.IGNORECASE)
    
    # Agravante/Agravado (Com suporte a plurais e hifens)
    full_text = re.sub(r'Agravante\s*\(?s?\)?\s*[:\-]', 'Agravante:', full_text, flags=re.IGNORECASE)
    full_text = re.sub(r'Agravad[oa]\s*\(?a?\)?\s*\(?s?\)?\s*[:\-]', 'Agravado:', full_text, flags=re.IGNORECASE)
    
    # Apelante/Apelado
    full_text = re.sub(r'Apelante\s*\(?s?\)?\s*[:\-]', 'Apelante:', full_text, flags=re.IGNORECASE)
    full_text = re.sub(r'Apelad[oa]\s*\(?a?\)?\s*\(?s?\)?\s*[:\-]', 'Apelado:', full_text, flags=re.IGNORECASE)
    
    terms_to_highlight = []
    
    # === STOP WORDS ===
    stop_words_list = [
        "Advogado", "Advogada", "OAB", "Juiz", "Relator", "Processo:", 
        "Intimação", "Data de", "DECISÃO", "SENTENÇA", "DESPACHO", "AGRAVO", "Vistos", 
        "Adv:", "Adv -", "Adv.", "MK_PARTES", "Parte intimação", "Parte Intimação"
    ]

    def extract_between_markers(start_marker, text_source):
        results = []
        text_upper = text_source.upper()
        marker_upper = start_marker.upper()
        
        start_idx = 0
        while True:
            idx = text_upper.find(marker_upper, start_idx)
            if idx == -1: break
            content_start = idx + len(marker_upper)
            
            min_stop = len(text_source)
            for sw in stop_words_list:
                sw_idx = text_upper.find(sw.upper(), content_start)
                if sw_idx != -1 and sw_idx > content_start and sw_idx < min_stop:
                    min_stop = sw_idx
            
            # Stop extra para OAB
            oab_match = re.search(r'[A-Z]{2}\s*-\s*\d{4,}', text_source[content_start:])
            if oab_match:
                abs_oab_idx = content_start + oab_match.start()
                if abs_oab_idx < min_stop:
                    min_stop = abs_oab_idx

            chunk = text_source[content_start:min_stop].strip()
            chunk = chunk.rstrip('.,-:')
            
            if chunk: results.append(chunk)
            start_idx = min_stop
            
        return results

    # A. Partes (MK_PARTES)
    partes_chunks = extract_between_markers("MK_PARTES:", full_text)
    for chunk in partes_chunks:
        if "SIGILO" in chunk.upper(): terms_to_highlight.append("SIGILO")
        corporate_shredder(chunk, terms_to_highlight)

    # B. Papéis Específicos
    roles = [
        "Recorrente", "Recorrido", "Interessado", "Impetrante", "Impetrado",
        "Exequente", "Executado", "Agravante", "Agravado", 
        "Autor", "Réu", "Apelante", "Apelado", 
        "Embargante", "Embargado", "Litisconsorte", "Representante", "Suscitante", "Suscitado",
        "Impugnante", "Impugnado", "Requerente", "Requerido"
    ]
    for role in roles:
        chunks = extract_between_markers(f"{role}:", full_text)
        for chunk in chunks:
            if is_valid_start(chunk):
                corporate_shredder(chunk, terms_to_highlight)

    # C. Processos
    regex_cnj = r"\d{7}[\s.-]?\d{2}[\s.]?\d{4}[\s.]?\d[\s.]?\d{2}[\s.]?\d{4}"
    regex_agravo = r"\d[\d\.]+\-\d\/\d{3,4}"
    regex_cnj_barra = r"\d{7}[\s.-]?\d{2}[\s.]?\d{4}[\s.]?\d[\s.]?\d{2}[\s.]?\d{4}\/\d+"
    
    processes = re.findall(regex_cnj_barra, full_text) + re.findall(regex_cnj, full_text) + re.findall(regex_agravo, full_text)
    terms_to_highlight.extend(processes)

    if not terms_to_highlight:
        return 0

    unique_terms = list(set(terms_to_highlight))
    return apply_highlight_reconstructor(paragraphs, unique_terms)

def corporate_shredder(text_block, target_list):
    if not text_block: return

    corp_suf = r'(LTDA|S[\.\/]?A\.?|EIRELI|LIMITADA|S\.S\.?|S\/C|ADVOCACIA|PARTICIPA[ÇC][ÕO]ES|FUNDO|INVESTIMENTO)'
    
    # 1. Quebra
    text_block = re.sub(rf'\b({corp_suf})(?:[\.\s,]*)(?:ME|EPP|EM\s+RECUPERA[ÇC][ÃA]O\s+JUDICIAL)?(?:[\.\s,]+|$)', r'\1 ### ', text_block, flags=re.IGNORECASE)
    
    # 2. CAIXA
    text_block = re.sub(r'(FEDERAL)\s+(?=[A-Z])', r'\1 ### ', text_block, flags=re.IGNORECASE)
    
    # 3. Barras
    text_block = re.sub(r'S/A', 'S_A_TEMP', text_block, flags=re.IGNORECASE)
    text_block = re.sub(r'S/C', 'S_C_TEMP', text_block, flags=re.IGNORECASE)
    text_block = text_block.replace('/', ' ### ')
    text_block = text_block.replace('S_A_TEMP', 'S/A').replace('S_C_TEMP', 'S/C')

    parts = re.split(r'(?:;|###)', text_block)
    
    for part in parts:
        p = part.strip().strip('.,;:-/¿? ')
        p_clean = re.sub(r'\s*[-–]?\s*(EM\s+)?RECUPERAÇÃO\s+JUDICIAL.*$', '', p, flags=re.IGNORECASE)
        p_clean = re.sub(r'\s*[-–]?\s*MASSA\s+FALIDA.*$', '', p_clean, flags=re.IGNORECASE)
        p_clean = re.sub(r'\s*\(Representado.*?\)', '', p_clean, flags=re.IGNORECASE)
        p_clean = re.sub(r'\s*\(A\)\s+.*', '', p_clean, flags=re.IGNORECASE)

        if validate_final_term(p_clean):
            target_list.append(p_clean)

# ==============================================================================
# 5. RECONSTRUTOR
# ==============================================================================
def apply_highlight_reconstructor(paragraphs, terms):
    count = 0
    terms = sorted(terms, key=len, reverse=True)
    
    for para in paragraphs:
        original_text = para.text
        if not original_text: continue
        
        # Sincronia
        clean_text_search = re.sub(r'[^\w\s\.\-\/\(\),:]', ' ', original_text)
        clean_text_search = re.sub(r'\s+', ' ', clean_text_search)
        
        keywords = ["Advogado", "Advogada", "OAB", "Relator", "Juiz", "Poder Judiciário", "Intimação", "Processo:", "Adv:", "Adv -"]
        for kw in keywords:
            clean_text_search = re.sub(rf'(\S)({kw})', r'\1 \2', clean_text_search, flags=re.IGNORECASE)
        
        clean_text_lower = clean_text_search.lower()
        
        found_in_para = [t for t in terms if t.lower() in clean_text_lower]
        if not found_in_para: continue

        highlights_map = []
        for term in found_in_para:
            start = 0
            while True:
                idx = clean_text_lower.find(term.lower(), start)
                if idx == -1: break
                
                is_overlap = False
                for h_start, h_end in highlights_map:
                    if (idx >= h_start and idx < h_end) or (idx + len(term) > h_start and idx + len(term) <= h_end):
                        is_overlap = True; break
                    if idx <= h_start and (idx + len(term)) >= h_end:
                        highlights_map.remove((h_start, h_end))
                        is_overlap = False; break
                
                if not is_overlap:
                    highlights_map.append((idx, idx + len(term)))
                    count += 1
                start = idx + 1
        
        if not highlights_map: continue
        highlights_map.sort()
        
        font_name = None; font_size = None
        try:
            if para.runs:
                font_name = para.runs[0].font.name
                font_size = para.runs[0].font.size
        except: pass

        para.clear()
        current_cursor = 0
        
        for start, end in highlights_map:
            if start > current_cursor:
                run = para.add_run(clean_text_search[current_cursor:start])
                if font_name: run.font.name = font_name
                if font_size: run.font.size = font_size
            
            run = para.add_run(clean_text_search[start:end])
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.bold = True
            if font_name: run.font.name = font_name
            if font_size: run.font.size = font_size
            
            current_cursor = end
            
        if current_cursor < len(clean_text_search):
            run = para.add_run(clean_text_search[current_cursor:])
            if font_name: run.font.name = font_name
            if font_size: run.font.size = font_size

    return count

@app.route('/progress')
def progress():
    return jsonify(current_status)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file')
        if not file: return jsonify({"error": "Envie um arquivo"}), 400

        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)

            update_status("Sanitizando XML...", 5)
            clean_filepath = sanitize_docx_xml(filepath)
            
            doc = Document(clean_filepath)
            update_status("Processando Publicações...", 10)
            
            count = process_document_by_publication(doc)
            
            update_status("Salvando...", 98)
            base_name = filename.replace('.docx', '')
            output_filename = f"GRIFADO_{base_name}_{int(time.time())}.docx"
            output_path = os.path.join(UPLOAD_FOLDER, output_filename)
            doc.save(output_path)
            
            update_status("Concluído!", 100)
            return jsonify({"status": "success", "download_url": f"/download/{output_filename}", "count": count})

        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    return render_template('index.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename), as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000)